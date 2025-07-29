import streamlit as st
import datetime
from openai import OpenAI
from docx import Document
from docx.shared import Pt

# Configuration: System prompt with detailed reporting guidelines
SYSTEM_PROMPT = st.secrets.get(
    "SYSTEM_PROMPT",
    """
You are an experienced disability examiner who writes Social Security Disability consultative exam reports following the exact DDS template:
- Use Times New Roman, 12 pt, black font
- Headers (e.g., 'History of Present Illness:') must be bold
- The header block (Name, SSN, DOB, Date of Exam) is single‑spaced, no extra spacing, followed by a blank line before Chief Complaint
- Chief Complaint line begins with 'Chief Complaint:' on the same line as its content
- Each report section should be a clear professional paragraph without extraneous commentary
"""
)

# Initialize OpenAI Client
client = OpenAI(api_key=st.secrets['OPENAI_API_KEY'])

def ai_generate(section_name, notes):
    prompt = f"""
You are an experienced disability examiner. Write a concise, professional paragraph for the '{section_name}' section of a DDS report based on these notes:

{notes}
"""
    try:
        response = client.chat.completions.create(
            model='gpt-3.5-turbo',
            messages=[
                {'role': 'system', 'content': SYSTEM_PROMPT},
                {'role': 'user', 'content': prompt}
            ],
            temperature=0.2,
            max_tokens=300
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        if hasattr(e, 'status') and e.status == 429:
            st.error('Rate limit exceeded. Please wait and try again.')
        else:
            st.error(f"Error generating {section_name}: {e}")
        return ''


def generate_report(data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Header block
    for field in ['name', 'ssn', 'dob', 'exam_date']:
        label = field.replace('_', ' ').title()
        p = doc.add_paragraph(f"{label}: {data.get(field, '')}")
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph('')  # blank line
    doc.add_paragraph(f"Chief Complaint: {data.get('chief_complaint', '')}")

    # Helper to add section
    def section(title, content):
        p = doc.add_paragraph()
        p.add_run(f"{title}:").bold = True
        doc.add_paragraph(content)

    # Report sections in specified order
    section('History of Present Illness - onset, diagnosis, treatment, medication, current status, pain severity, conditional items', data.get('hpi', ''))
    section('Past Medical History - list of past conditions', data.get('pmh', ''))
    section('Family History - list of conditions known to run in family', data.get('family', ''))
    section('Social History - relationship status, living situation, children, employment status, support from family or friends, alcohol, tobacco or drug use', data.get('social', ''))
    section('Review of Systems - current reported symptoms', data.get('ros', ''))
    section('Activities of Daily Living - Wake up, Bed Time, Work status, Chores, Daily Life', data.get('adl', ''))
    # Physical Examination composite
    phys = (
        f"General Observations: {data.get('general_observations', '')}\n\n"
        f"Vitals: {data.get('vitals', '')}\n\n"
        f"HEENT: {data.get('heent', '')}\n\n"
        f"Respiratory: {data.get('respiratory', '')}\n\n"
        f"Cardiac: {data.get('cardiac', '')}\n\n"
        f"Digestive: {data.get('digestive', '')}\n\n"
        f"Upper Extremities: {data.get('upper_extremities', '')}\n\n"
        f"Lower Extremities: {data.get('lower_extremities', '')}\n\n"
        f"Cervical and Thoracic Spines: {data.get('spines', '')}"
    )
    section('Physical Examination', phys)
    section('Neuro', data.get('neuro', ''))
    section('Skin', data.get('skin', ''))
    section('Assessment and Impressions', data.get('assessment', ''))
    section('Conclusion – Functional Ability', data.get('capacity', ''))

    filename = f"{data.get('name', '').replace(' ', '_')}_DDS_Report.docx"
    doc.save(filename)
    return filename

# Streamlit UI
st.title('DDS Exam Report Builder with GPT Assistance')

# Input order
data = {}
# Basic demographics
data['name'] = st.text_input('Name')
data['ssn'] = st.text_input('SSN (Last 4)')
data['dob'] = st.date_input('Date of Birth', min_value=datetime.date(1900,1,1), max_value=datetime.date.today())
data['exam_date'] = st.date_input('Date of Exam')
# Chief Complaint
data['chief_complaint'] = st.text_area('Chief Complaint - Provided by DDS')

# Sections requiring AI drafting
sections = [
    ('hpi', 'History of Present Illness - onset, diagnosis, treatment, medication, current status, pain severity, conditional items'),
    ('pmh', 'Past Medical History - list of past conditions'),
    ('family', 'Family History - list of conditions known to run in family'),
    ('social', 'Social History - relationship status, living situation, children, employment status, support from family or friends, alcohol, tobacco or drug use'),
    ('ros', 'Review of Systems - current reported symptoms'),
    ('adl', 'Activities of Daily Living - Wake up, Bed Time, Work status, Chores, Daily Life'),
]
# Initialize and render AI sections
for key, label in sections:
    st.session_state.setdefault(f'{key}_notes', '')
    st.session_state.setdefault(key, '')
    st.text_area(f'{label} Notes', key=f'{key}_notes')
    if st.button(f'Draft {label}', key=f'btn_{key}'):
        st.session_state[key] = ai_generate(label, st.session_state[f'{key}_notes'])
    data[key] = st.text_area(label, key=key)

# Physical exam inputs
st.subheader('Physical Examination Details')
phy_fields = [
    ('general_observations','General Observations'),
    ('vitals','Vitals - BP, Pulse, Height, Weight, Vision (corrected or non-corrected) Right - Left - Both'),
    ('heent','HEENT'),
    ('respiratory','Respiratory'),
    ('cardiac','Cardiac'),
    ('digestive','Digestive'),
    ('upper_extremities','Upper Extremities'),
    ('lower_extremities','Lower Extremities'),
    ('spines','Cervical and Thoracic Spines'),
]
for key, label in phy_fields:
    st.session_state.setdefault(f'{key}_notes', '')
    st.session_state.setdefault(key, '')
    st.text_area(f'{label} Notes', key=f'{key}_notes')
    if st.button(f'Draft {label}', key=f'btn_{key}'):
        st.session_state[key] = ai_generate(label, st.session_state[f'{key}_notes'])
    data[key] = st.text_area(label, key=key)

# Remaining sections
remaining = [
    ('neuro', 'Neuro'),
    ('skin', 'Skin'),
    ('assessment', 'Assessment and Impressions'),
    ('capacity', 'Conclusion – Functional Ability')
]
for key, label in remaining:
    st.session_state.setdefault(f'{key}_notes', '')
    st.session_state.setdefault(key, '')
    st.text_area(f'{label} Notes', key=f'{key}_notes')
    if st.button(f'Draft {label}', key=f'btn_{key}'):
        st.session_state[key] = ai_generate(label, st.session_state[f'{key}_notes'])
    data[key] = st.text_area(label, key=key)

# Generate report button
if st.button('Generate Word Document'):
    file_path = generate_report(data)
    with open(file_path, 'rb') as f:
        st.download_button('Download DDS Report', f, file_name=file_path)
